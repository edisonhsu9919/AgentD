"""DOCX extraction layer (Phase O2).

Pure data extraction from DOCX files using python-docx.
No LLM calls — structured raw data only.
"""

import os
from typing import Any

# Text sample limits
_TEXT_SAMPLE_MAX_PARAGRAPHS = 10
_TEXT_SAMPLE_MAX_CHARS = 2000
_HEADING_MAX = 20


def extract(path: str) -> dict[str, Any]:
    """Extract structural metadata and text sample from a DOCX file.

    Returns a dict with:
      - path, kind, office_kind, size_bytes
      - paragraph_count, heading_count, table_count
      - headings (list of heading texts, capped)
      - text_sample (first N paragraphs, truncated)
      - metadata (title, author, subject)

    Raises FileNotFoundError if path doesn't exist.
    Raises ValueError if file is not a valid DOCX.
    """
    if not os.path.isfile(path):
        raise FileNotFoundError(f"File not found: {path}")

    size_bytes = os.path.getsize(path)

    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx") from e

    try:
        doc = Document(path)
    except Exception as e:
        raise ValueError(f"Cannot read DOCX: {e}") from e

    # Extract metadata
    props = doc.core_properties
    metadata = {
        "title": _clean(props.title),
        "author": _clean(props.author),
        "subject": _clean(props.subject),
    }

    # Collect paragraphs and headings
    paragraphs: list[str] = []
    headings: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            headings.append(text)

    # Count tables
    table_count = len(doc.tables)

    # Build text sample from first N paragraphs
    text_sample = _build_text_sample(paragraphs)

    return {
        "path": path,
        "kind": "office",
        "office_kind": "docx",
        "size_bytes": size_bytes,
        "paragraph_count": len(paragraphs),
        "heading_count": len(headings),
        "table_count": table_count,
        "headings": headings[:_HEADING_MAX],
        "text_sample": text_sample,
        "metadata": metadata,
    }


def _build_text_sample(paragraphs: list[str]) -> str:
    """Build a text sample from the first few paragraphs, capped."""
    parts: list[str] = []
    total = 0
    for para in paragraphs[:_TEXT_SAMPLE_MAX_PARAGRAPHS]:
        remaining = _TEXT_SAMPLE_MAX_CHARS - total
        if remaining <= 0:
            break
        chunk = para[:remaining]
        parts.append(chunk)
        total += len(chunk)
    return "\n\n".join(parts)


def _clean(value) -> str:
    if value is None:
        return ""
    return str(value).strip()
